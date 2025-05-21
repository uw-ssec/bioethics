from __future__ import annotations

import re
from pathlib import Path

import markdown2
import requests
import streamlit as st  # type: ignore[import]
import urllib3
from bs4 import BeautifulSoup
from config import (
    API_BASE_URL,
    EMBEDDING_MODEL,
    EXISTING_COLLECTION,
    EXISTING_QDRANT_PATH,
    GENERATION_MODEL,
    PROMPT_TEMPLATES,
)
from docx import Document
from langchain.document_loaders import PyMuPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter

urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

CHUNK_SIZE = 1200
CHUNK_OVERLAP = 120

text_splitter = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)

filter_keywords = [
    "references",
    "acknowledgements",
    "author contributions",
    "bibliography",
    "funding",
]

# Initialize session state for messages if not exists
if "messages" not in st.session_state:
    st.session_state.messages = []

st.title("Bioethics RAG Chatbot")

template_options = ["No Template", *list(PROMPT_TEMPLATES.keys())]
selected_template = st.selectbox("Choose a prompt template", template_options, index=0)


def is_noise_chunk(chunk: str) -> bool:
    noise_patterns = [
        r"^\s*©",
        r"copyright",
        r"ceur workshop",
        r"issn",
        r"arxiv",
        r"http[s]?://",
        r"license",
        r"all rights reserved",
    ]
    chunk_lower = chunk.lower()
    return any(re.search(pat, chunk_lower) for pat in noise_patterns)


def process_uploaded_files(uploaded_files):
    """Reads uploaded PDF files and extracts text + metadata."""
    documents = []

    for file in uploaded_files:
        with Path(file.name).open("wb") as f:
            f.write(file.getbuffer())
        loader = PyMuPDFLoader(file.name)  # Load PDF
        pages = loader.load()

        for page in pages:
            page_content_lower = page.page_content.lower()
            skip_page = False
            for keyword in filter_keywords:
                if (
                    keyword in page_content_lower
                    and (
                        page.page_content.count(keyword) > 1
                        or any(
                            line.strip().lower().startswith(keyword)
                            for line in page.page_content.split("\n")
                        )
                    )
                    and re.search(
                        r"^\s*(" + keyword + r")\s*$\n",
                        page.page_content,
                        re.IGNORECASE | re.MULTILINE,
                    )
                ):
                    skip_page = True
                    break
            if skip_page:
                continue

            chunks = text_splitter.split_text(page.page_content)
            for i, chunk in enumerate(chunks):
                if is_noise_chunk(chunk):
                    continue
                documents.append({"page_content": chunk, "metadata": {**page.metadata, "chunk": i}})

    return documents


def retrieve_response(query, documents):
    """
    Calls the backend retrieval API to fetch relevant documents based on the query.

    Parameters:
        query (str): The user's question or search query
        documents (list): List of PDF dicts with page_content and metadata

    Returns:
        dict: JSON response from the API containing:
            - docs: List of retrieved document chunks
            - status_code: HTTP status code of the response
    """
    try:
        response = requests.post(
            f"{API_BASE_URL}/retrieve/",
            json={
                "query": query,
                "documents": documents,
                "embedding_model": EMBEDDING_MODEL,
                "existing_collection": EXISTING_COLLECTION,
                "existing_qdrant_path": EXISTING_QDRANT_PATH,
            },
        )
        response.raise_for_status()
        return response.json()
    except requests.RequestException as e:
        st.error(f"❌ Retrieval API failed: {e!s}")
        return {"docs": []}


def generate_response(prompt):
    """
    Calls the backend generation API to produce a generated answer to the prompt.

    Parameters:
        prompt (str): a formatted prompt ready to send to the generation model
        generation_model (str): The model to use for text generation

    Returns:
        dict: JSON response from the API containing:
            - answer: The generated text response
            - status_code: HTTP status code of the response
    """
    try:
        response = requests.post(
            f"{API_BASE_URL}/generate/",
            json={"prompt": prompt, "generation_model": GENERATION_MODEL},
        )
        response.raise_for_status()
        return response.json()
    except requests.RequestException as e:
        st.error(f"❌ Generation API failed: {e!s}")
        return {"answer": "⚠️ Failed to generate response."}


def save_to_docx(content: str, filename: str = "output.docx"):
    """
    Saves the generated content to a .docx file.

    Parameters:
        content (str): The text content to save.
        filename (str): The name of the .docx file.
    """
    doc = Document()
    doc.add_paragraph(content)
    doc.save(filename)


def markdown_to_docx_via_html(markdown_content: str, filename: str = "output.docx"):
    """
    Converts Markdown content to a .docx file via HTML.

    Parameters:
        markdown_content (str): The Markdown content to convert.
        filename (str): The name of the .docx file.
    """
    # Convert Markdown to HTML
    html_content = markdown2.markdown(markdown_content)

    # Parse HTML and add content to a Word document
    doc = Document()
    soup = BeautifulSoup(html_content, "html.parser")
    for element in soup.descendants:
        if element.name == "p":
            doc.add_paragraph(element.get_text())
        elif element.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            doc.add_heading(element.get_text(), level=int(element.name[1]))
        elif element.name == "li":
            doc.add_paragraph(element.get_text(), style="List Bullet")
        elif element.name == "table":
            # Handle tables
            table = doc.add_table(rows=0, cols=0)
            for row in element.find_all("tr"):
                cells = row.find_all(["td", "th"])
                doc_row = table.add_row()
                for i, cell in enumerate(cells):
                    if len(doc_row.cells) <= i:
                        table.add_column()
                    doc_row.cells[i].text = cell.get_text()

    # Save the document
    doc.save(filename)


if selected_template == "No Template":
    uploaded_files = st.file_uploader(
        "Attach documents (PDFs)", type=["pdf"], accept_multiple_files=True
    )
    documents = process_uploaded_files(uploaded_files) if uploaded_files else []
    if query := st.chat_input("Your question:"):
        st.session_state.messages.append({"role": "user", "content": query})
        with st.chat_message("user"):
            st.markdown(query)
        retrieved_docs = []
        retrieved_text = ""
        if documents or EXISTING_COLLECTION or EXISTING_QDRANT_PATH:
            with st.spinner("Retrieving relevant documents..."):
                retrieved_docs = retrieve_response(query, documents).get("docs", [])
                retrieved_text = "\n\n".join(
                    doc["page_content"] for doc in retrieved_docs if doc.get("page_content")
                )
            if retrieved_docs:
                with st.chat_message("assistant"):
                    st.markdown("### Retrieved Document Chunks:")
                    for doc in retrieved_docs:
                        st.markdown(f"- {doc['page_content'][:500]}")
        with st.spinner("Generating response..."):
            implicit_template = "Context:\n{context}\n\nQuestion: {question}"
            formatted_prompt = implicit_template.format(context=retrieved_text, question=query)
            generate_response_data = generate_response(formatted_prompt)
            if "answer" in generate_response_data:
                generated_answer = generate_response_data["answer"]
                markdown_to_docx_via_html(generated_answer, filename="generated_output.docx")
                st.success("✅ Response saved as 'generated_output.docx'")
                with Path("generated_output.docx").open("rb") as f:
                    docx_bytes = f.read()
                st.download_button(
                    label="Download .docx",
                    data=docx_bytes,
                    file_name="generated_output.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            else:
                st.error("⚠️ Unable to generate a response. Please try again later.")
                generated_answer = "⚠️ Failed to generate response."
        assistant_message = {
            "role": "assistant",
            "content": generated_answer,
            "chunks": [doc["page_content"][:500] for doc in retrieved_docs],
        }
        st.session_state.messages.append(assistant_message)
        with st.chat_message("assistant"):
            st.markdown(generated_answer)
else:
    user_template = st.text_area(
        "Customize template", value=PROMPT_TEMPLATES[selected_template], height=250
    )
    uploaded_files = st.file_uploader(
        "Attach documents (PDFs)", type=["pdf"], accept_multiple_files=True
    )
    documents = process_uploaded_files(uploaded_files) if uploaded_files else []
    if st.button("Generate Summary/Report"):
        st.session_state.messages.append(
            {
                "role": "user",
                "content": f"[{selected_template}] Generate summary/report for uploaded documents.",
            }
        )
        with st.chat_message("user"):
            st.markdown(f"[{selected_template}] Generate summary/report for uploaded documents.")
        retrieved_docs = []
        retrieved_text = ""
        if documents or EXISTING_COLLECTION or EXISTING_QDRANT_PATH:
            with st.spinner("Retrieving relevant documents..."):
                retrieved_docs = retrieve_response("", documents).get("docs", [])
                retrieved_text = "\n\n".join(
                    doc["page_content"] for doc in retrieved_docs if doc.get("page_content")
                )
            if retrieved_docs:
                with st.chat_message("assistant"):
                    st.markdown("### Retrieved Document Chunks:")
                    for doc in retrieved_docs:
                        st.markdown(f"- {doc['page_content'][:500]}")
        with st.spinner("Generating response..."):
            formatted_prompt = user_template.format(context=retrieved_text)
            generate_response_data = generate_response(formatted_prompt)
            if "answer" in generate_response_data:
                generated_answer = generate_response_data["answer"]
                markdown_to_docx_via_html(generated_answer, filename="generated_output.docx")
                st.success("✅ Response saved as 'generated_output.docx'")
                with Path("generated_output.docx").open("rb") as f:
                    docx_bytes = f.read()
                st.download_button(
                    label="Download .docx",
                    data=docx_bytes,
                    file_name="generated_output.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            else:
                st.error("⚠️ Unable to generate a response. Please try again later.")
                generated_answer = "⚠️ Failed to generate response."
        assistant_message = {
            "role": "assistant",
            "content": generated_answer,
            "chunks": [doc["page_content"][:500] for doc in retrieved_docs],
        }
        st.session_state.messages.append(assistant_message)
        with st.chat_message("assistant"):
            st.markdown(generated_answer)
