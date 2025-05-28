from __future__ import annotations

import markdown2
from bs4 import BeautifulSoup
from config import UW_PURPLE
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def add_uw_header(doc: Document) -> None:
    header = doc.sections[0].header
    header.is_linked_to_previous = False  # Ensure header is unique for this section
    # Clear any existing content in the header (e.g., default paragraph)
    for p in header.paragraphs:
        p.clear()
    for t in header.tables:
        header._element.remove(t._element)

    htable = header.add_table(rows=1, cols=2, width=Inches(6))

    # Left cell
    left_cell = htable.cell(0, 0)
    left_paragraph = left_cell.paragraphs[0] if left_cell.paragraphs else left_cell.add_paragraph()
    run = left_paragraph.add_run()
    run.add_picture("frontend/uw_logo.png", width=Inches(0.75))
    left_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Right cell
    right_cell = htable.cell(0, 1)
    right_paragraph = (
        right_cell.paragraphs[0] if right_cell.paragraphs else right_cell.add_paragraph()
    )
    run = right_paragraph.add_run("University of Washington")
    run.font.name = "Arial"
    run.font.size = Pt(10)
    right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def save_to_docx(markdown_content: str, filename: str = "generated_output.docx"):
    """
    Saves the generated markdown content to a .docx file via HTML.

    Parameters:
        markdown_content (str): The Markdown content to convert.
        filename (str): The name of the .docx file.
    """
    html_content = markdown2.markdown(markdown_content)
    doc = Document()
    add_uw_header(doc)

    soup = BeautifulSoup(html_content, "html.parser")
    for element in soup.descendants:
        if element.name == "p":
            children = list(element.children)
            if (
                len(children) == 1
                and getattr(children[0], "name", None) == "strong"
                and element.get_text(strip=True) == children[0].get_text(strip=True)
            ):
                heading = doc.add_heading(element.get_text(strip=True), level=3)
                for run_item in heading.runs:
                    run_item.font.color.rgb = UW_PURPLE
                    run_item.font.name = "Arial"
            else:
                doc.add_paragraph(element.get_text())
        elif element.name == "li":
            doc.add_paragraph(element.get_text(), style="List Bullet")
        elif element.name == "table":
            table = doc.add_table(rows=0, cols=0)
            for row in element.find_all("tr"):
                cells = row.find_all(["td", "th"])
                doc_row = table.add_row()
                for i, cell in enumerate(cells):
                    if len(doc_row.cells) <= i:
                        table.add_column()
                    doc_row.cells[i].text = cell.get_text()

    doc.save(filename)
