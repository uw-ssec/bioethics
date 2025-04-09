import streamlit as st
import requests
import json
import time
import textwrap
from io import BytesIO
from docx import Document
from langchain.document_loaders import PyMuPDFLoader
from config import (
    API_BASE_URL,
    GENERATION_MODEL
)

st.title("Document Summarization")

default_template = textwrap.dedent("""
    You are a medical science expert and have to write a report on the below papers:
    
    {context}
    
    Write a summary to communicate the research to study participants in a few sentences for each section. Write in lay language at a 6th grade reading level.

    Headings for the summary:
    - What was the research about?
    - How was the research done?
    - What did the researchers learn? (Answer this in detailed bullet points)
    - What was new and innovative about the studies?
    - What do the findings mean?
    - What's next?
""")

# Let the user view/edit the prompt template
user_template = st.text_area("Edit Prompt Template", value=default_template, height=250)

# Prompt template function
def summary_prompt(context: str) -> str:
    """Format the retrieval context into the final prompt."""
    return textwrap.dedent(f"""
    You are a medical science expert and have to write a report on the below papers:
    
    {context}
    
    Write a summary to communicate the research to study participants in a few sentences for each section. Write in lay language at a 6th grade reading level.

    Headings for the summary:
    - What was the research about?
    - How was the research done?
    - What did the researchers learn?(Answer this in detailed bullet points)
    - What was new and innovative about the studies?
    - What do the findings mean?
    - What's next?
                           
    """)

def custom_summary_prompt(context: str, template: str) -> str:
    return template.replace("{context}", context)


# Initialize session state for messages if not already set
if "messages" not in st.session_state:
    st.session_state.messages = []

# Display chat history (if any)
for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])

# File uploader for documents
uploaded_files = st.file_uploader("Attach documents (PDFs)", type=["pdf"], accept_multiple_files=True)

# Function to process uploaded PDFs
def process_uploaded_files(uploaded_files):
    """Reads uploaded PDF files and extracts text + metadata."""
    documents = []
    for file in uploaded_files:
        # Save the file temporarily so PyMuPDFLoader can load it
        with open(file.name, "wb") as f:
            f.write(file.getbuffer())
        loader = PyMuPDFLoader(file.name)
        pages = loader.load()
        for page in pages:
            documents.append({
                "page_content": page.page_content,
                "metadata": page.metadata
            })
    return documents

# Process PDFs if any are uploaded
documents = process_uploaded_files(uploaded_files) if uploaded_files else []

if documents:
    # Combine all document pages into a single context string
    full_context = "\n\n".join(doc["page_content"] for doc in documents)

    # Button to trigger summarization
    def generate_summary():
        # Log the user action
        st.session_state.messages.append({"role": "user", "content": "Summarize the uploaded documents."})
        with st.chat_message("user"):
            st.markdown("Summarize the uploaded documents.")

        # Format the prompt using the full document context
        prompt = custom_summary_prompt(full_context, user_template)
        
        with st.spinner("Generating summary..."):
            try:
                generate_payload = {
                    "prompt": prompt,
                    "generation_model": GENERATION_MODEL
                }
                generate_response = requests.post(f"{API_BASE_URL}/generate/", json=generate_payload)
                if generate_response.status_code == 200:
                    generated_answer = generate_response.json().get("answer", "")
                else:
                    generated_answer = "⚠️ Failed to generate summary."
            except Exception as e:
                generated_answer = "⚠️ Failed to generate summary."
                st.error(f"❌ Generation API failed: {str(e)}")

        # Display the generated summary
        assistant_message = {
            "role": "assistant",
            "content": generated_answer
        }
        st.session_state.messages.append(assistant_message)

        with st.chat_message("assistant"):
            st.markdown(generated_answer)
    
        st.session_state.generated_summary = generated_answer
    
    if st.button("Summarize"):
        generate_summary()
    
    if "generated_summary" in st.session_state and st.button("Regenerate"):
        generate_summary()
        
    if "generated_summary" in st.session_state:
        # Create a DOCX document and add the generated summary
        doc = Document()
        doc.add_heading("Document Summary", level=1)
        doc.add_paragraph(st.session_state.generated_summary)
        # Save document to a BytesIO stream
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        st.download_button(
            label="Save Output to DOCX",
            data=doc_io,
            file_name="document_summary.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
else:
    st.info("Please upload one or more PDF documents to generate a summary.")
